from __future__ import annotations

import re
from pathlib import Path


def inspect_document(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return inspect_pdf(path)
    if suffix == ".docx":
        return inspect_docx(path)
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return summarize_text(path, text, kind=suffix.lstrip("."))
    raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")


def inspect_pdf(path: Path) -> dict:
    pages, full_text = extract_pdf_pages(path)

    return {
        "filename": path.name,
        "kind": "pdf",
        "page_count": len(pages),
        "characters": sum(page["characters"] for page in pages),
        "preview": compact_preview(full_text, limit=900),
        "pages": pages[:12],
    }


def extract_pdf_pages(path: Path, max_chars: int | None = None) -> tuple[list[dict], str]:
    import pdfplumber

    pages = []
    full_text = []
    characters = 0
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            characters += len(text)
            if max_chars is None or len("\n".join(full_text)) < max_chars:
                full_text.append(text)
            pages.append(
                {
                    "page": index,
                    "characters": len(text),
                    "preview": compact_preview(text),
                }
            )
    joined = "\n".join(full_text)
    if max_chars is not None:
        joined = joined[:max_chars]
    for page in pages:
        page.setdefault("characters", 0)
    if pages and sum(page["characters"] for page in pages) != characters:
        pages[0]["characters"] += characters - sum(page["characters"] for page in pages)
    return pages, joined


def paper_brief_from_pdf(path: Path) -> dict:
    if path.suffix.lower() != ".pdf":
        raise ValueError("Only PDF papers can be imported into the research inquiry cards.")
    pages, text = extract_pdf_pages(path, max_chars=32000)
    compact_text = compact_preview(text, limit=24000)
    if not compact_text:
        raise ValueError("No selectable text was found in this PDF. Use an OCR/text PDF before importing.")
    lines = clean_document_lines(text)
    title = infer_paper_title(lines, path.stem)
    abstract = extract_named_section(
        text,
        starts=("abstract", "摘要"),
        ends=("keywords", "key words", "index terms", "introduction", "1 introduction", "引言", "关键词"),
    )
    keywords = extract_keywords(text)
    focus_source = abstract or compact_preview("\n".join(lines[:16]), limit=700)
    tensions = extract_tension_sentences(text)
    brief = {
        "start_mode": "research",
        "topic": title[:600],
        "research_focus": compact_preview(focus_source, limit=900),
        "assumptions": [
            "以导入论文的研究问题、证据范围与方法设定作为真实研究依据",
            f"PDF 文本层包含 {len(pages)} 页、约 {sum(page['characters'] for page in pages)} 个字符",
        ],
        "stakeholders": infer_stakeholders(text),
        "tensions": tensions,
        "input_sources": {
            "topic": "research",
            "research_focus": "research",
            "assumptions": "research",
            "stakeholders": "research",
            "tensions": "research",
        },
        "import_document": {
            "filename": path.name,
            "kind": "pdf",
            "page_count": len(pages),
            "characters": sum(page["characters"] for page in pages),
            "preview": compact_preview(text, limit=900),
            "keywords": keywords,
        },
    }
    if keywords:
        brief["assumptions"].append("论文关键词：" + "；".join(keywords[:6]))
    return {
        "filename": path.name,
        "kind": "pdf",
        "page_count": len(pages),
        "characters": sum(page["characters"] for page in pages),
        "preview": compact_preview(text, limit=900),
        "pages": pages[:12],
        "paper_brief": brief,
    }


def clean_document_lines(text: str) -> list[str]:
    lines = []
    for raw in str(text or "").splitlines():
        line = " ".join(raw.split()).strip()
        if not line:
            continue
        if re.fullmatch(r"\d+", line):
            continue
        lines.append(line)
    return lines


def infer_paper_title(lines: list[str], fallback: str) -> str:
    blocked = re.compile(r"^(abstract|摘要|keywords?|key words|doi|arxiv|copyright|proceedings|journal|vol\.|volume|\d+\.)\b", re.I)
    for line in lines[:40]:
        if blocked.search(line):
            continue
        if 8 <= len(line) <= 220 and len(line.split()) <= 32:
            return line
    return fallback.replace("_", " ").replace("-", " ").strip() or "Imported paper"


def extract_named_section(text: str, starts: tuple[str, ...], ends: tuple[str, ...]) -> str:
    compact = re.sub(r"\s+", " ", str(text or " ")).strip()
    if not compact:
        return ""
    start_pattern = "|".join(re.escape(item) for item in starts)
    end_pattern = "|".join(re.escape(item) for item in ends)
    match = re.search(rf"(?:^|\s)({start_pattern})[:：]?\s*(.+?)(?=(?:^|\s)({end_pattern})[:：]?|$)", compact, flags=re.I)
    if not match:
        return ""
    return compact_preview(match.group(2), limit=1200)


def extract_keywords(text: str) -> list[str]:
    section = extract_named_section(
        text,
        starts=("keywords", "key words", "index terms", "关键词"),
        ends=("introduction", "1 introduction", "引言", "background", "背景"),
    )
    parts = re.split(r"[,;；、|]", section)
    keywords = []
    for item in parts:
        clean = item.strip(" .:：-")
        if 2 <= len(clean) <= 48 and clean.lower() not in {"keywords", "key words", "index terms"} and clean not in keywords:
            keywords.append(clean)
    return keywords[:8]


def extract_tension_sentences(text: str) -> list[str]:
    pattern = re.compile(r"(uncertain|uncertainty|risk|challenge|limitation|trade[- ]off|bias|safety|privacy|伦理|风险|不确定|挑战|限制|偏见|安全|隐私)")
    sentences = re.split(r"(?<=[。.!?？])\s+", re.sub(r"\s+", " ", str(text or "")))
    result = []
    for sentence in sentences:
        clean = sentence.strip()
        if 18 <= len(clean) <= 220 and pattern.search(clean.lower()) and clean not in result:
            result.append(clean)
        if len(result) >= 3:
            break
    return result or ["论文结论进入未来情境时，研究证据、实际部署与受影响人群之间仍存在不确定性"]


def infer_stakeholders(text: str) -> list[str]:
    lower = str(text or "").lower()
    candidates = []
    mapping = [
        (("patient", "patients", "clinical", "hospital", "患者", "医院", "临床"), "患者、医生与医疗机构"),
        (("student", "teacher", "education", "learning", "学生", "教师", "教育"), "学习者、教师与教育机构"),
        (("worker", "employee", "labor", "workplace", "工人", "员工", "劳动"), "一线使用者、管理者与组织决策者"),
        (("user", "participant", "consumer", "用户", "参与者", "消费者"), "系统使用者与研究参与者"),
        (("policy", "government", "regulation", "政策", "监管", "政府"), "政策制定者与监管者"),
    ]
    for needles, label in mapping:
        if any(needle in lower for needle in needles) and label not in candidates:
            candidates.append(label)
    if not candidates:
        candidates.extend(["论文作者/研究团队", "被研究对象或实验参与者", "未来系统使用者"])
    return candidates[:5]


def inspect_docx(path: Path) -> dict:
    from docx import Document

    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        tables.append(
            {
                "table": table_index,
                "rows": len(table.rows),
                "columns": len(table.columns),
            }
        )

    text = "\n".join(paragraphs)
    return {
        "filename": path.name,
        "kind": "docx",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "characters": len(text),
        "preview": compact_preview(text, limit=900),
        "tables": tables[:12],
    }


def summarize_text(path: Path, text: str, kind: str) -> dict:
    return {
        "filename": path.name,
        "kind": kind,
        "line_count": len(text.splitlines()),
        "characters": len(text),
        "preview": compact_preview(text, limit=900),
    }


def compact_preview(text: str, limit: int = 220) -> str:
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return compact[: limit - 1].rstrip() + "…"
